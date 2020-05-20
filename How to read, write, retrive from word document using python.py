"""
Document to write para in word doc, read para from list
"""
from docx import Document

class Doc:
    def write_para(self, *para):
        """
        Write n number of paragraph into word document
        """        
        document = Document()
        for i in para:
            paragraph = document.add_paragraph(i)
            document.save('test1.docx')

    def all_para(self):
        """
        Get all paragraph in a list.
        """
        para_list = []
        doc = Document('test1.docx')
        all_paras = doc.paragraphs
        for para in all_paras:
            para_list.append(para.text)
        #print("Total paragraph in the document is '", len(para_list), "'")
        return para_list
        
    def get_para(self, para_no):
        """
        Retrieve pragraph content based on the list index.
        """
        try:
            para_list = Doc.all_para(self)
            print("Paragraph content of ", para_no, "=>")
            print("       ",para_list[para_no])
        except:
            print("Invalid para index.")
b1 = Doc()
para1 = "To the Participants and Beneficiaries of the Amerisafe, Inc. 401(k) Plan."
para2 = "As permitted by 29 CFR 2520.103-8 of the Department of Labor’s Rules."
para3 = "Task to write into doc."
para4 = "Because of the significance of the information."
para5 = "Because of the significance of the information that I did not audit."
b1.write_para(para1, para2, para3, para4, para5)
m = b1.all_para()
print(m)
print("****************************************************")
b1.get_para(4)

